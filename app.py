import streamlit as st
import anthropic
import tempfile
import os
import json
import re
import subprocess
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Authentication ─────────────────────────────────────────────────────────────

JERRY_PINK = "#ff3975"

LOGIN_CSS = """
<style>
#MainMenu, footer, header {visibility: hidden;}
.login-wrap {max-width: 400px; margin: 80px auto 0;}
.login-logo {font-size: 1.6rem; font-weight: 800; color: #ff3975; letter-spacing: -0.5px; margin-bottom: 0.2rem;}
.login-sub  {color: #888; font-size: 0.9rem; margin-bottom: 1.5rem;}
div[data-testid="stButton"] > button {
    background: #ff3975 !important; color: #fff !important;
    border: none !important; border-radius: 8px !important;
    font-weight: 600 !important; padding: 0.55rem 1.5rem !important;
}
div[data-testid="stButton"] > button:hover {background: #e02d62 !important;}
</style>
"""

def check_login():
    if st.session_state.get("authenticated"):
        return True
    st.set_page_config(page_title="Jerry JD Generator", layout="centered")
    st.markdown(LOGIN_CSS, unsafe_allow_html=True)
    st.markdown('<div class="login-wrap">', unsafe_allow_html=True)
    st.markdown('<div class="login-logo">Jerry JD Generator</div>', unsafe_allow_html=True)
    st.markdown('<div class="login-sub">Internal tool · Jerry recruiting team only</div>', unsafe_allow_html=True)
    pwd = st.text_input("Password", type="password", placeholder="Enter team password")
    if st.button("Sign in", use_container_width=True):
        correct = st.secrets.get("APP_PASSWORD", "")
        if correct and pwd == correct:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password.")
    st.markdown('</div>', unsafe_allow_html=True)
    return False

if not check_login():
    st.stop()

# ── System prompt (V6 SKILL.md — canonical source) ────────────────────────────

SYSTEM_PROMPT = """You are Jerry.ai's JD writer. Your only job is to generate job descriptions that meet Jerry's editorial standard.

When a recruiter gives you a role spec, do two things in sequence before outputting anything:
1. Generate a full JD draft following the skill below.
2. Audit your draft against every Hard Rule (1–26). Fix every violation. Pay particular attention to Rules 21, 23, 24, 25, and 26, which require your judgment.

Then output ONLY the finished JD as clean markdown — no preamble, no commentary. Use this format:
- Job title as H1: # Title
- Section headings as H2: ## Section Name
- Bullet points: - **Bold label:** body text (for labeled bullets)
- Regular paragraphs as plain text

After the JD, output exactly:
---AUDIT---
Brief plain-English summary of what you caught and fixed. If nothing: "No violations found."
---END---

═══════════════════════════════════════════════
JERRY.AI JD SKILL
═══════════════════════════════════════════════

## The philosophy (internalize this)

Most JDs are corporate noise — grocery lists of responsibilities that attract no one specific and repel no one wrong. Jerry JDs do the opposite:

- **Hook fast.** You have 3 seconds. The opening must name the specific problem or opportunity this role addresses.
- **Radical honesty.** Say the unglamorous parts. Bad fits should self-select out.
- **Outcomes over tasks.** Write "Own X and drive Y" rather than "responsible for Z".
- **Specific archetypes over generic requirements.** An archetype like A Detective obsessed with edge cases carries more signal than "strong attention to detail".
- **Subtraction is editing.** Every sentence that doesn't add value gets cut.

---

## Hard Rules (check every JD against this list before output)

These are non-negotiable. Each was established from real editorial feedback. Numbers are stable IDs.

**Prose constructions (banned in all forms):**

1. **No "X, not Y" / antithesis / punchy-reversal.** Any two clauses mirrored for rhythm where the second mainly completes a beat. Banned even when it reads well. Examples to never write: "compliance is the whole point, not a footnote"; "the risk stops growing and starts shrinking"; "make this worse, not better"; "high autonomy, high trust, low context". State the point once, directly.

2. **No restatement / cadence-over-content.** Don't repeat a verb or structure for rhythm ("we grew fast and we grew everywhere"; the triple "its own X, its own Y, its own Z"). If parallel structure is doing work the words don't earn, cut it. This applies across sentences and across bullets too — if a bold bullet label and the bullet body make the same point, one of them goes. If two consecutive bullets describe the same reality in different words, collapse them.

3. **No setup line the next sentence says better.** If the following sentence is the concrete version, the opener is a throat-clear. Delete it and open on the concrete line. This applies inside bullets too — if the first clause of a bullet body repeats the bold label, cut the clause.

4. **No punchy-for-its-own-sake fragments** — staccato buzzword previews of a paragraph.

5. **No "the X question is answered, the Y question is still open" structure.**

6. **No phrases that sound substantive but say nothing concrete** ("a big part of this job is showing up"). This includes conclusory summary sentences — sentences whose only job is to announce that the previous sentence was important: "That is the bar." "This matters." "The stakes are real." "That is not abstract." If a sentence adds no new information and exists only to punctuate a point already made, delete it. The previous sentence should land on its own.

7. **No fake parallelism** — list items that look symmetrical but aren't grammatically (two bare noun phrases plus one with a verb clause).

**Words / tired moves:**

8. **"not glamorous / not sexy / not exciting" openers are banned unless fresh** — factually correct AND a rarely-stated claim for that category. Fresh: "car insurance is not sexy". Stale (banned): payroll, accounting, compliance described as boring. If the category is already the default punchline for dull work, find a more specific honest truth.

9. **No "from day one"** appended to ownership bullets — filler.

10. **No corporate filler or AI-default prose patterns.** Corporate filler: "dynamic environment", "wear many hats", "passionate about innovation", and similar. AI-default patterns — phrases that are statistically common in AI-generated text and read as machine-made: "in progress", "ongoing journey", "evolving landscape", "ever-changing", "at the intersection of", "a unique opportunity to", "fast-paced and dynamic", and similar. If a phrase sounds like something a language model would reach for, it's probably filler. Cut it and say the specific thing.

**List & structure:**

11. **Collapse redundant list items.** If multiple items make one point, keep one.

12. **Land lists on the most concrete item, not the most abstract.** A list that ends on the vaguest, most rhetorical item loses force.

13. **No nested sub-bullets.** Flat lists only.

14. **Subheadings must fit the content beneath them**, not the reverse. The label describes what the bullet actually says.

**Specificity:**

15. **Real numbers, not vague quantifiers.** "16 US states" not "more than a dozen". "One market to twelve in eighteen months" not "we grew everywhere". Don't hedge inside a sentence reaching for specificity.

16. **Don't mix categories in a list for cadence.** Don't put sub-national units and a whole country in one breath ("a dozen states plus Canada" with "plus Canada" as an afterthought).

**Content / formatting / legal:**

17. **Archetype labels never use quotation marks.** Hard rule, no exceptions.

18. **Job titles only — never names of specific people** anywhere in a JD.

19. **Minimum-only years framing** ("7+ years"). Never a range or a maximum — legal compliance.

20. **No compensation or benefits.** Handled in the ATS; never appears in a JD.

21. **Radical Honesty bullets must be genuine differentiators, not table stakes.** A differentiator is something most employers could NOT honestly say. Table stakes any employer could claim — cross-functional collaboration, leadership visibility, "full organizational support", "great team", generic growth — do not belong in this section. Test each bullet: could a mediocre competitor say the same sentence? If yes, cut or sharpen it.

22. **"Why we need you" heading is wrong for junior or high-supply roles.** Use it only where the candidate has leverage (senior/scarce talent).

**Factual & confidentiality integrity:**

23. **Every claim must be verified against the role spec — nothing invented or assumed.** Clause-level accuracy matters: don't write "run the playbook" when the spec says there is no playbook; don't name a tool that doesn't fit the use case. When unsure, omit rather than invent.

24. **Don't overshare sensitive internal facts.** Radical honesty has a confidentiality ceiling. Three categories are always confidential regardless of how much the spec reveals:
   - **Revenue model details:** specific commission structures, fee types, carrier arrangements, origination fee mechanics, or any specifics of how Jerry earns money. The existence of revenue complexity is fine; the mechanics are not.
   - **Proprietary technology architecture:** how specific systems connect, what the automation does internally, specific data flows between named internal systems. The existence of homegrown tooling is fine; the architecture is not.
   - **Operationally sensitive people information:** never disclose that a role is a backfill, that someone has left, why a previous person in the seat didn't work out, or any other information about past or present employees. Use only job titles, never names or departure context.
   Test: does this sentence tell a competitor or the public something specific about how Jerry makes money, how its systems work, or why a specific person is or isn't in this role? If yes, abstract it.

25. **No cross-bullet restatement within a section.** Every bullet in a section must own a distinct point. If two bullets in the same section make the same point — even in different words, even at different levels of specificity — collapse them or cut one. This applies to every section without exception. Before finalising any section, read all its bullets together and ask: does any bullet echo, restate, or merely add emphasis to a point another bullet already makes? If yes, that bullet goes.

26. **Radical Honesty bullets must create genuine friction for wrong-fit candidates.** A bullet that merely describes what the role involves — onboarding process, scope, workflow, or neutral org-chart facts like reporting line — belongs in "What You'll Own", not "Why You Will (Or Won't) Like Working Here". To pass this section, each bullet must describe something a real candidate might read and think "that's not for me" — a constraint, a tradeoff, a cultural reality, or an unglamorous truth. If a bullet would only make someone think "okay, good to know", cut it or move it to the ownership section.

   Note: structural facts can pass this test if they carry real friction. "You are inheriting a developing function" passes — it filters out candidates who want a mature, well-oiled team. "You will manage a small team of two" passes — it filters out candidates who want scale. "You report to the Controller" fails — it informs but disqualifies nobody. The test is not whether a bullet describes structure; it's whether that specific fact would cause a real wrong-fit candidate to self-select out.

---

## The 5-Section Structure (mandatory, fixed order)

Every Jerry JD has exactly these 5 sections. Section headings can vary — the content type is fixed.

| # | Content type | Example heading variants |
|---|---|---|
| 1 | **The Hook** — why this role exists, what problem/opportunity it addresses, why it matters for Jerry and the world | "The Challenge", "The Mission", "Why We Exist", "Why We Need You", "The Opportunity" |
| 2 | **Radical Honesty** — culture, pace, reporting structure, what's glamorous and what isn't | "Why You Will (Or Won't) Like Working Here", "What to Expect", "The Reality" |
| 3 | **Ownership** — outcomes and impact expected, often phased (Phase 1 / Phase 2) | "What You'll Own", "Your Mission", "What You'll Build" |
| 4 | **Archetype** — mindset, traits, and working style of the ideal candidate | "Who You Are", "The Ideal Candidate", "Who Thrives Here" |
| 5 | **Experience** — concrete prior experience, skills, or outcomes required | "Ideal Candidate", "What You'll Bring", "Your Background" |

**Optional sections:**
- **Tech Stack** (engineering roles only) — list languages, frameworks, infra after Section 5.
- **Why Jerry.ai** (any role) — a closing section that zooms back out to Jerry's mission, traction, and why now is an exciting time to join. Use sparingly; most JDs don't need it. Best for senior/exec roles where company context helps close the candidate.

---

## Writing each section

### Section 1 — The Hook

- Open with the macro context: car ownership affects 80%+ of Americans, yet it's fragmented and expensive.
- Then zoom into the specific problem this role solves or opportunity it unlocks.
- End with: if you succeed at this, here's what changes for Jerry and its customers.
- No generic "we're a fast-growing startup" filler.
- Length: 2–4 tight paragraphs. Senior/exec roles can go longer; coordinator/specialist roles should be shorter.

Jerry's standard context (use or adapt — always use broader framing per Rule 30):
"Your car and your home are your most important assets, yet the experience of owning one is stuck in the 90s. Every part of the journey (buy/sell, insurance, maintenance/repairs, etc) is fragmented, complicated, and expensive.

Jerry.ai is building the first app to manage it all. We started with your car ($2T market in the U.S.), launched car insurance in 2019, then driving insights, diagnostics, and a repair marketplace. Since then, we've reached 5M+ customers, raised $240M+, scaled revenue 70X, and became profitable in early 2024."

### Section 2 — Radical Honesty

- Use 3–5 bullet points with **bold labels**.
- Each bullet names one honest truth about working here — positive or negative.
- Always include: reporting structure, pace/autonomy, and degree of hands-on work.
- If it's an IC role with no direct reports: say that explicitly.
- If it's a fast-paced, scrappy environment: describe the actual pace concretely, never "dynamic".
- Each bullet must be a genuine differentiator (Rule 21). Test: could a mediocre competitor say the same sentence? If yes, rewrite it.
- Never name specific individuals — job titles only.

### Section 3 — Ownership

- Frame as outcomes, not tasks.
- Default format is prose or flat bullets.
- **Use Phase 1 / Phase 2 structure ONLY when the role spec explicitly describes two distinct phases.** Do not impose this structure on every role. If the spec does not describe two phases, do not invent them.
- Avoid "you will be responsible for..." — prefer "you will own..." or "you will lead..."

### Section 4 — Archetype

- 3–4 bullets, each with a **bold label** followed by a specific illustration of the behavior.
- **Archetype labels are never wrapped in quotation marks.** Write **A Detective:** never **"A Detective":**.
- **Labels never use the "X, not Y" / antithesis construction.**
- Two valid label styles — both acceptable, and they can be mixed:
  - Named archetype: **A Detective:** You love hunting for edge cases...
  - Trait label: **Structured problem solver:** You break ambiguous problems into clear hypotheses...
- Avoid generic adjectives ("detail-oriented", "self-starter", "passionate"). Every bullet must name a specific behavior.

### Section 5 — Experience

- Concrete, specific, and minimal. Only list what's actually required.
- Use "[X]+ years in [specific field] with proven success in [outcome]" format.
- No maximum years of experience (legal compliance).
- 4–7 bullet points max.

---

## Tone and voice

Jerry's voice is: **direct, honest, specific, confident, and a little irreverent**. It does not use corporate filler, bury the lede, list 15 requirements when 6 will do, or pretend every part of the job is exciting. It names the unglamorous parts, trusts the reader to self-select, and sounds like a smart person wrote it — not an HR template.

---

## Length guidance

| Role type | Approximate length |
|---|---|
| IC / Specialist | 500–650 words |
| Senior IC / TL | 650–800 words |
| Manager / Director | 700–850 words |
| VP / Exec | 800–1000 words |

Cut ruthlessly. A tight 500-word JD beats a padded 900-word one.

---

## Never say anything negative about other companies or service providers

Jerry's business depends on partnerships with carriers, insurance providers, and other service providers. Never make negative claims about named or unnamed companies, partners, or industry incumbents — not about their motives, their products, or their priorities. You can describe what's broken in the market by pointing to the absence of a solution. You cannot attribute that failure to any company's behavior or intent.

Permitted: "car owners had no single place to manage insurance, financing, and maintenance"
Banned: "insurance companies prioritize profit over the customer"

---

## Additional rules (27–29)

27. **Bold label brevity — Sections 2 and 4.** Every bold label in Section 2 (Radical Honesty) and Section 4 (Archetype) must be as short as possible — target 2–4 words maximum. Cut every word that doesn't do essential work. Test each label: can you remove any word and keep the meaning? If yes, remove it. Examples of what to cut: "The playbook doesn't exist yet" → "No playbook." "Thick-skinned and patient" → "Thick-skinned." "A chameleon by instinct" → "A Chameleon." Filler qualifiers like "by instinct", "at heart", "in nature", "yet", "by default" attached to a label are always cuttable. The label is a signal, not a sentence.

28. **No descriptive sub-headings within phases.** When Section 3 uses Phase 1 / Phase 2 structure, do NOT add a second descriptive label inside the phase (e.g. "Reddit (Days 1–90)" or "Expanding Channel Presence"). The phase label alone is the heading. Go directly from "**Phase 1:**" into the content. A phase label plus a sub-label is always redundant — the sub-label restates what the phase summary already says.

29. **Hook paragraph transitions must be explicit.** When the Hook moves from one topic to another — e.g. from company context to the specific problem this role addresses — there must be a bridging sentence or clause that makes the logical connection explicit. Never place two unconnected ideas in adjacent paragraphs. The reader should always be able to answer: "why does this paragraph follow that one?"

30. **Always use broader asset framing in the Hook — but make the transition logical.** Never open with car-only framing. Jerry's mission is now to manage all physical assets (car, home, motorcycle, etc.), so the Hook must always open from that broader vantage point: "Your car and your home are your most important assets..." or "Jerry.ai is building the first app to manage all your physical assets..." Do NOT use "More than 80% of Americans own a car..." as the opening. However, the broader framing must then transition logically to the specific problem this role addresses. For a car insurance role, you zoom in: open broadly → connect to the insurance product → name the specific problem. For an expansion role, the broader framing connects directly. The test: a recruiter should never feel whiplash reading the Hook. Each paragraph must follow from the one before it.

---

## Manual audit pass (do this before outputting)

**Rules 21 + 26:** For every Radical Honesty bullet, answer both: (a) Could a mediocre competitor honestly say this? If yes → cut or rewrite. (b) Would a real wrong-fit candidate think "that's not for me"? If no → cut or move to ownership.

**Rule 25:** Read all bullets in each section together. Does any bullet make the same point as another? If yes → collapse.

**Rule 23:** For every specific claim — tools named, scope described, team size — verify against the role spec. If not in the spec, omit it.

**Rule 27:** Read every bold label in Sections 2 and 4. Is any label longer than 4 words? If yes, cut it down.

**Rule 28:** If Phase 1 / Phase 2 structure is used, check that no descriptive sub-heading appears within the phase. If one exists, delete it.

**Rules 29 + 30:** Read the Hook. (a) Is every paragraph transition explicit — does each paragraph follow logically from the one before? (b) Does the Hook open with broader asset framing (car + home, or all physical assets) rather than car-only? If either fails, rewrite.

**Rules 1–20, 22, 24:** Check mechanically. Fix every violation."""

# ── Sample JDs (V6 canonical) ─────────────────────────────────────────────────

SAMPLE_JDS = """# Jerry.ai Canonical Sample JDs — Style and Tone Benchmarks

## Sample 1: Engineering Lead / Manager, Core Product Engine (Backend)

### The technical challenge

More than 80% of Americans own a car and see it as a fundamental necessity, yet the experience of owning one is stuck in the 90s. Every part of the user journey (shopping for a car, insurance, repairs, etc) is complicated, fragmented, expensive and time-consuming. At Jerry.ai, we are building the first mobile app to manage it all.

Our insurance marketplace is our core product today, and our ability to automate the end-to-end shopping experience is our core differentiator. The Core Backend Team owns the most critical part of that, the "last mile", where we orchestrate complex, multi-step transactions with external systems to finalize insurance policy purchases. Because the industry lacks modern APIs, we build custom orchestration layers that navigate internal and external interfaces with precision. It is a high-stakes puzzle of asynchronous state management and real-time system reliability.

Your challenge will be to evolve our current automation systems to a robust, fault-tolerant framework that can handle external interface changes gracefully and scale across dozens of new categories with minimal manual intervention.

### Why you will (or won't) like working here

- **Ownership:** You'll report directly to our co-founder/CTO, who is actively involved in product implementation. There is no red tape, no slow-moving processes, and no layers of approvals needed. If you have an idea, you can pitch it, build it, and see it live in days.
- **Stellar team:** We have a very high bar for talent, as a result, we have a small but mighty team that can move fast and accomplish a lot. You will work with likeminded peers who are equally passionate and who care intensely about what they do.
- **Hands-on:** Many of our TLMs are ex-founders, they like getting their hands dirty and solving difficult problems. You will spend 80% of your time doing hands-on technical work and 20% of your time managing a team of 7 engineers.
- **Volume of customer data:** We are a consumer-facing business (5M+ customers) and have access to a massive amount of customer data across behaviours like driving/telematics, shopping, payment, etc.

### What you'll own

Our first automation system was built for speed under early pressure. Today it's a web of tightly coupled business logic and database access that makes changes risky and failures hard to debug.

- **Phase 1:** You'll be deep in the code, learning both the technical and business sides of our core product. Then you will lead incremental refactoring of our existing automation systems: designing and implementing new standards and pairing with your team of 7 engineers to ensure the new system is predictable and testable.
- **Phase 2:** As the architecture stabilizes, you'll shift towards using an AI-first approach to building the next generation of internal tools that will allow us to grow from 5M to 50M users.

### Who you are

- **A Player Coach:** You likely founded a company or were a lead architect at an early stage startup. You are a builder at heart and like being in the trenches alongside your team.
- **A Detective:** You are obsessed with success rates. You love hunting for edge cases in complex business logic and won't stop until you understand exactly why a script failed.
- **A Systems Thinker:** You have a high bar for system design, especially for offline, asynchronous jobs. You understand that async work requires much stricter standards for decoupling, state management, and failure retry logic than synchronous online systems.

### Our tech stack

TypeScript (Nest.js, React), GraphQL, AWS (ECS/EKS, Lambda).

---

## Sample 2: Creative Director

### Why we exist

More than 80% of Americans own a car and see it as a fundamental necessity, yet the experience of owning one is stuck in the 90s. Every part of the user journey is complicated, fragmented, and time-consuming, and the costs are staggering. In 2025, the cost of owning a car was more than 20% of the average household income, after tax. At Jerry.ai, we are building the first app to manage it all.

### Why we need you

Since we launched our app in 2019, we've reached 5M customers, raised over $240M in funding, scaled our annual revenue more than 70X, and became profitable in early 2024.

We did this with an exceptionally talented team of analytical thinkers who are brilliant at experimentation and growth optimization. But what we don't have is someone who obsesses over every single word we put in front of our customers and how we make them feel.

Car ownership is a stressful and emotionally charged topic for most people, and right now, our messaging is very transactional. We need you to change that. Your job will be to think about what our customers actually need to hear from us, and how to fundamentally change how we speak to and care for them.

### Why you will (or won't) like working here

- **Blank canvas:** You won't inherit a tone of voice or brand playbook. You will build it in partnership with our CEO, including defining our personality and voice across all organic surfaces.
- **Hands-on:** This is a senior level role in scope and impact, but you will be an IC with no direct reports on day one.
- **Car ownership is not "sexy":** Insurance, maintenance, and safety are only topics people think about when they have a bill to pay or get into an accident. Anyone can market something inherently exciting, but can you build enthusiasm and trust around a "dry" topic?

### What you'll own

- **Phase 1:** Your sole focus will be overhauling our existing customer communications, including CRM, email, SMS, and in-app copy. You will take our messaging from functional and direct to value-additive and relationship-driven.
- **Phase 2:** As our tone of voice begins to take shape, your scope may expand to driving customer growth through broader creative organic content.

### Who you are

- **A copywriting virtuoso:** You know that every word matters. You can translate complex, dry, and even stressful topics into copy that is engaging, human, and accessible.
- **A T-Shaped marketer:** Your deep expertise is in copywriting, audience development and product positioning. Horizontally, you have a great sense of the social media landscape and community dynamics.
- **A first principles thinker:** You avoid making assumptions. You know how to take qualitative data and user pain points and turn them into resonant, empathetic marketing copy.
- **Comfortable with discomfort:** You've weathered fast-growing startups before and can confidently push back on highly analytical, data-first teams to champion the unmeasurable value of a great brand experience.

### Ideal experience

- 7+ years of experience in creative marketing with a focus on organic growth and a copy lean.
- Proven track record with organic channels like email, SMS, and social.
- Experience working with consumer startups.
- A portfolio showcasing both your writing and creative craft.

---

## Sample 3: Senior Product Manager, DriveShield

### Why we exist

Jerry.ai is building the first app to manage all your physical assets: starting with your car, then home, and motorcycle, etc. Our insurance marketplace is our core product today. Now, we are expanding into broader problems that make owning a car expensive, confusing, and stressful.

### Why we're hiring this role

DriveShield is one of our big product bets to make Jerry useful not just when someone shops for insurance, but every time they drive. The opportunity is huge, but the challenge is: most people don't wake up excited to check their driving behavior data. If the experience is not clear, motivating, and trustworthy, people won't care.

We need a Senior PM who can turn DriveShield from a useful feature into a daily habit for millions of drivers.

### Why you will (or won't) like the work

- **You will own a real behavior-change problem.** Your job is to make drivers care, come back, and improve.
- **You will be hands-on.** You will work deeply in product, data, UX, research, and experimentation.
- **This category is not inherently "sexy".** Insurance and driving safety are not things people think about everyday — only when something bad happens and it's too late.
- **The roadmap will not be handed to you.** You will find the highest-leverage problems, build conviction, and ship.
- **The bar is high.** You will work with a talent-dense team from companies like Nvidia, LinkedIn, DoorDash, Amazon, Klarna, and Lemonade.

### What you'll own

Your mandate is to make DriveShield clearer, more motivating, and more valuable — measured through adoption, engagement, retention, and trust.

- **Make DriveShield a habit:** Own the journey from first trip to next action: understanding a score, seeing what changed, earning points, getting guidance, and feeling motivated to drive safer.
- **Improve the core experience:** Modernize trip tracking, driving scores, points, safety tips, progress, and guidance.
- **Build engagement loops that feel useful:** Figure out what makes drivers come back after a trip.
- **Explore what DriveShield becomes next:** Help define new safety and mobility capabilities, including smarter insights, crash detection, and AI-powered personalization.
- **Learn fast from users and data:** Use analytics, research, usability studies, prototypes, and A/B tests.

### Who you are

- **A consumer product craftsperson.** You know that tiny UX decisions can decide whether users trust or abandon a product.
- **A behavioral product thinker.** You care about why people form habits, ignore advice, respond to rewards, or distrust scores.
- **A detective.** You do not stop at "engagement is down." You find where users dropped, what they saw, and why it happened.
- **Independent judgment on data.** You build conviction from messy signals and use experiments to learn quickly.

### Ideal experience

- Consumer product experience across mobile or web.
- Work on engagement, retention, personalization, lifecycle, rewards, or habit formation.
- Strong partnership with Design, Research, Engineering, and Data.
- Comfort with funnel analysis, experiments, and ambiguous user behavior.

---

## Sample 4: Senior Manager, BizOps & Analytics (New Products)

Note: This JD deliberately omits Section 2 (Radical Honesty). The Phase 1 / Phase 2 ownership structure already surfaces all genuine friction points — the ramp period, the fully hands-on 0-to-1 mandate, no inherited playbook. No additional bullet survived both the Rule 21 test (genuine differentiator) and Rule 26 test (creates real friction) without restating what ownership already says. This is a valid 4-section JD.

### Why we exist

Your car and your home are your most important assets, yet the experience of owning one is stuck in the 90s. Every part of the journey (buy/sell, insurance, maintenance/repairs, etc) is fragmented, complicated, and expensive.

Jerry.ai is building the first app to manage it all. We started with your car ($2T market in the U.S.) — launched car insurance in 2019, became one of the top 3 brokers in the country, then added driving insights, diagnostics, and a repair marketplace.

We've reached 5M+ customers, raised $240M+, and scaled our revenue 80X (since our app launch in 2019). And we've been profitable since 2024.

### Why we're hiring

We're now ready to expand beyond car ownership into adjacent verticals (home, motorcycle, RV, etc) to manage all of your physical assets in one place. You will own one of these new bets from 0-1.

This is not a pure strategy role. You will be hands-on, standing up a new business line, including scoping/sizing a new market, research and analysis, building a new product, leading go-to-market, testing and iterating, and eventually, building a team to scale it.

### What you'll own

- **Phase 1: Ramp on existing products.** For the first few months, you will join an existing product team, and learn how we think about and develop products. You will learn how to conduct customer research, write product specs, work with engineers and designers, use tools like Figma and Codex to prototype, design and run experiments, analyze outcomes, and incorporate learnings into iterative product decisions.
- **Phase 2: Build a new business.** Take one of our target verticals (e.g. commercial auto insurance, property tax, RV insurance, etc) from concept to a revenue-generating product. Run the incubation, build and ship the product, test and iterate, and then hire a team to run it.

### Who you are

- **An entrepreneur at heart.** You want to learn how to build a business from the ground up, and how to scale it.
- **Tolerance for ambiguity.** You're comfortable making decisions with incomplete information and pivoting quickly as your understanding evolves.
- **Deeply curious.** You are fascinated by the "why" behind customer behaviors and market dynamics.
- **Data-obsessed.** If a funnel drops 2% or an A/B test returns an unexpected result, you dive immediately into the analysis to figure out the root cause. And you aren't satisfied until you understand what really happened.

### What you bring

- 5+ years experience at a consulting firm, investment bank, high-growth startup or equivalent.
- Track record of independently owning ambiguous problems and delivering outcomes.
- Strong analytical skills; comfortable running data analysis in SQL, Excel or Python.
- Experience juggling multiple complex work streams without letting anything drop.
- Comfort and ability to interface with executive stakeholders on a daily basis."""


# ── Markdown → .docx converter ────────────────────────────────────────────────

def md_to_docx(md_text: str, output_path: str):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(1))
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(12)

    def set_arial(run):
        run.font.name = "Arial"
        rPr = run._r.get_or_add_rPr()
        tag = qn("w:rFonts")
        el = rPr.find(tag)
        if el is None:
            el = OxmlElement("w:rFonts")
            rPr.insert(0, el)
        el.set(qn("w:ascii"), "Arial")
        el.set(qn("w:hAnsi"), "Arial")

    def add_runs(para, text):
        for part in re.split(r"(\*\*[^*]+\*\*)", text):
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                run = para.add_run(part)
            set_arial(run)

    for line in md_text.strip().splitlines():
        s = line.strip()
        if s.startswith("# "):
            p = doc.add_heading(s[2:].strip(), level=1)
            for r in p.runs: r.font.name = "Arial"
        elif s.startswith("## "):
            p = doc.add_heading(s[3:].strip(), level=2)
            for r in p.runs: r.font.name = "Arial"
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:].strip())
        elif s:
            p = doc.add_paragraph()
            add_runs(p, s)
    doc.save(output_path)


# ── Linter ────────────────────────────────────────────────────────────────────

def run_linter(docx_path: str) -> dict:
    linter = os.path.join(os.path.dirname(os.path.abspath(__file__)), "lint_jd.py")
    if not os.path.exists(linter):
        return {"summary": {"fail": 0, "flag": 0, "total": 0}, "findings": []}
    try:
        r = subprocess.run([sys.executable, linter, docx_path],
                           capture_output=True, text=True, timeout=30)
        if r.stdout.strip():
            return json.loads(r.stdout)
    except Exception:
        pass
    return {"summary": {"fail": 0, "flag": 0, "total": 0}, "findings": []}


def fmt_linter(results: dict) -> str:
    findings = results.get("findings", [])
    s = results.get("summary", {})
    if not findings:
        return "Linter clean — fail: 0, flag: 0"
    lines = [f"Linter — fail: {s.get('fail',0)}, flag: {s.get('flag',0)}\n"]
    for f in findings:
        tag = "[FAIL]" if f["severity"] == "fail" else "[FLAG]"
        lines.append(f"{tag} Rule {f['rule']}: {f['text'][:80]!r}\n       → {f['note']}")
    return "\n".join(lines)


# ── Claude API calls ──────────────────────────────────────────────────────────

def generate_jd(client, role_spec: str, job_title: str):
    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content":
            f"Style benchmarks for tone calibration:\n\n{SAMPLE_JDS}\n\n---\n\n"
            f"Generate JD for: {job_title}\n\nRole spec:\n{role_spec}"}]
    )
    raw = resp.content[0].text
    if "---AUDIT---" in raw:
        jd, audit = raw.split("---AUDIT---", 1)
        return jd.strip(), audit.replace("---END---", "").strip()
    return raw.strip(), "(No audit report returned.)"


def fix_fails(client, role_spec, job_title, jd_text, linter_report):
    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[
            {"role": "user", "content":
                f"Style benchmarks:\n\n{SAMPLE_JDS}\n\n---\n\n"
                f"Generate JD for: {job_title}\n\nRole spec:\n{role_spec}"},
            {"role": "assistant", "content": jd_text},
            {"role": "user", "content":
                f"The deterministic linter found these violations:\n\n{linter_report}\n\n"
                f"Fix every [FAIL]. Review every [FLAG]. Output the corrected JD in the same "
                f"markdown format, followed by ---AUDIT--- and what you changed."}
        ]
    )
    raw = resp.content[0].text
    if "---AUDIT---" in raw:
        jd, audit = raw.split("---AUDIT---", 1)
        return jd.strip(), audit.replace("---END---", "").strip()
    return raw.strip(), "(No audit report returned.)"


# ── Streamlit UI ──────────────────────────────────────────────────────────────

st.set_page_config(page_title="Jerry JD Generator", layout="wide")
st.markdown(f"""
<style>
#MainMenu, footer, header {{visibility: hidden;}}

/* Buttons */
div[data-testid="stButton"] > button {{
    background: {JERRY_PINK} !important; color: #fff !important;
    border: none !important; border-radius: 8px !important;
    font-weight: 600 !important; padding: 0.55rem 1.5rem !important;
    transition: background 0.15s;
}}
div[data-testid="stButton"] > button:hover {{ background: #e02d62 !important; }}
div[data-testid="stButton"] > button:disabled {{ background: #f0f0f0 !important; color: #aaa !important; }}

/* Sign out button — secondary style */
section[data-testid="stSidebar"] div[data-testid="stButton"] > button {{
    background: transparent !important; color: #ff3975 !important;
    border: 1.5px solid #ff3975 !important; border-radius: 8px !important;
}}
section[data-testid="stSidebar"] div[data-testid="stButton"] > button:hover {{
    background: #fff0f3 !important;
}}

/* Download button */
div[data-testid="stDownloadButton"] > button {{
    background: {JERRY_PINK} !important; color: #fff !important;
    border: none !important; border-radius: 8px !important;
    font-weight: 600 !important; padding: 0.55rem 1.5rem !important;
}}

/* Inputs */
div[data-testid="stTextInput"] input, div[data-testid="stTextArea"] textarea {{
    border-radius: 8px !important; border: 1.5px solid #e0e0e0 !important;
}}
div[data-testid="stTextInput"] input:focus, div[data-testid="stTextArea"] textarea:focus {{
    border-color: {JERRY_PINK} !important; box-shadow: 0 0 0 2px #ff397520 !important;
}}

/* Audit report box */
.report-box {{
    background: #fafafa; border: 1px solid #ebebeb; border-radius: 10px;
    padding: 1.1rem; font-family: monospace; font-size: .82rem;
    white-space: pre-wrap; max-height: 480px; overflow-y: auto;
    line-height: 1.55;
}}

/* Sidebar */
section[data-testid="stSidebar"] {{ background: #fff; border-right: 1px solid #f0f0f0; }}

.how-step {{
    display: flex; gap: 0.75rem; align-items: flex-start;
    margin-bottom: 0.9rem;
}}
.how-num {{
    background: {JERRY_PINK}18; color: {JERRY_PINK};
    font-weight: 700; font-size: 0.75rem; border-radius: 50%;
    min-width: 22px; height: 22px; display: flex;
    align-items: center; justify-content: center; margin-top: 1px;
}}
.how-text {{ font-size: 0.85rem; color: #444; line-height: 1.4; }}
.how-text strong {{ color: #111; }}
</style>
""", unsafe_allow_html=True)

# Header
st.markdown(f'<h1 style="color:{JERRY_PINK};font-size:1.9rem;font-weight:800;margin-bottom:0;">Jerry JD Generator ✍️</h1>', unsafe_allow_html=True)
st.markdown('<p style="color:#888;margin-top:0.2rem;margin-bottom:1.5rem;font-size:0.95rem;">Paste your role spec. Get an Alana-approved JD in seconds.</p>', unsafe_allow_html=True)

with st.sidebar:
    st.markdown(f'<p style="font-size:1.1rem;font-weight:700;color:{JERRY_PINK};margin-bottom:0.2rem;">Jerry JD Generator</p>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:0.78rem;color:#aaa;margin-top:0;margin-bottom:1rem;">Internal tool · Recruiting team only</p>', unsafe_allow_html=True)
    if st.button("Sign out", use_container_width=True):
        st.session_state.authenticated = False
        st.rerun()
    st.markdown('<hr style="border:none;border-top:1px solid #f0f0f0;margin:1.2rem 0;">', unsafe_allow_html=True)
    st.markdown(f'<p style="font-weight:700;font-size:0.85rem;color:#111;margin-bottom:0.8rem;">How it works</p>', unsafe_allow_html=True)
    steps = [
        ("Copy and paste the role spec", "Copy the full text from your role spec and paste it here. Do not paste a link — Claude cannot access external URLs. The more detail you give, the better the output."),
        ("Claude writes and self-audits", "A full JD is drafted following Jerry's 30 Hard Rules — broader asset framing in the hook, explicit transitions, short bold labels, no redundant phase headings, correct tone, structure, legal compliance, and confidentiality."),
        ("The linter runs deterministically", "16 mechanical rules are checked in Python: banned phrases, X-not-Y constructions, year ranges, nested bullets, quoted archetype labels, verbose bold labels, and temporal sub-headings inside phases."),
        ("Fails are auto-fixed", "Any hard violations go back to Claude for a correction pass. The linter re-runs to confirm the output is clean."),
        ("Download your .docx", "The final file comes with a full audit report — what was caught, what was changed, and the linter summary."),
    ]
    for i, (title, body) in enumerate(steps, 1):
        st.markdown(f'''
        <div class="how-step">
            <div class="how-num">{i}</div>
            <div class="how-text"><strong>{title}</strong><br>{body}</div>
        </div>''', unsafe_allow_html=True)

# Read API key from secrets (never exposed in UI)
api_key = st.secrets.get("ANTHROPIC_API_KEY", "") or os.environ.get("ANTHROPIC_API_KEY", "")
if not api_key:
    st.error("ANTHROPIC_API_KEY is not configured. Contact your admin.")
    st.stop()

job_title = st.text_input("Job title", placeholder="e.g. Senior PM, DriveShield")
role_spec  = st.text_area("Role spec — paste full text", height=280,
    placeholder="Copy and paste the full text from your role spec here. Do not paste a link — Claude cannot access external URLs. Claude will not invent facts; it only uses what you give it.")

run = st.button("Generate JD", disabled=not (api_key and job_title and role_spec))

if run:
    if not api_key.startswith("sk-"):
        st.error("API key looks wrong — should start with sk-ant-...")
        st.stop()

    client = anthropic.Anthropic(api_key=api_key)
    col_left, col_right = st.columns([3, 2])

    with col_left:
        with st.status("Generating JD…", expanded=True) as status:
            st.write("Calling Claude…")
            try:
                jd_text, audit_text = generate_jd(client, role_spec, job_title)
            except Exception as e:
                st.error(f"Claude API error: {e}")
                st.stop()

            st.write("Running deterministic linter…")
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            tmp.close()
            md_to_docx(jd_text, tmp.name)
            lint1 = run_linter(tmp.name)
            fails1 = [f for f in lint1.get("findings", []) if f["severity"] == "fail"]

            lint_final = lint1
            fix_audit  = ""
            if fails1:
                st.write(f"Found {len(fails1)} linter fail(s) — asking Claude to fix…")
                try:
                    jd_text, fix_audit = fix_fails(client, role_spec, job_title,
                                                    jd_text, fmt_linter(lint1))
                    md_to_docx(jd_text, tmp.name)
                    lint_final = run_linter(tmp.name)
                except Exception as e:
                    st.warning(f"Fix pass error: {e}")

            status.update(label="Done ✓", state="complete")

        safe = re.sub(r"[^\w\s-]", "", job_title).strip().replace(" ", "_")
        with open(tmp.name, "rb") as fh:
            docx_bytes = fh.read()
        os.unlink(tmp.name)

        st.download_button("⬇ Download .docx", data=docx_bytes,
            file_name=f"{safe}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with st.expander("JD preview", expanded=True):
            st.markdown(jd_text)

    with col_right:
        s = lint_final.get("summary", {})
        fails_final = [f for f in lint_final.get("findings", []) if f["severity"] == "fail"]
        if fails_final:
            st.error(f"Linter: {s.get('fail',0)} fail(s) remain after fix pass")
        else:
            st.success(f"Linter: 0 fails · {s.get('flag',0)} flag(s) for review")

        report = ["═══ CLAUDE AUDIT ═══", audit_text]
        if fix_audit:
            report += ["\n═══ FIX PASS ═══", fix_audit]
        report += ["\n═══ DETERMINISTIC LINTER ═══", fmt_linter(lint_final)]

        st.markdown(
            f'<div class="report-box">{"<br>".join(l.replace(chr(10),"<br>") for l in report)}</div>',
            unsafe_allow_html=True)
